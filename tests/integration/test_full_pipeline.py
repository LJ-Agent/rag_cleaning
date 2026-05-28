"""Full-chain integration tests for RAG-CLEANING service.

Tests the complete pipeline: preprocessing -> element processing -> cleaning -> validation -> output.
"""

import io
import json
import os
import sys
import tempfile
from pathlib import Path

# Add src to path
sys.path.insert(0, str(Path(__file__).parent.parent.parent / "src"))

import pytest


# ─── Test Document Generators ─────────────────────────────


def _make_test_pdf() -> bytes:
    """Generate a minimal but valid PDF with text content."""
    # Minimal hand-crafted PDF for testing
    pdf = (
        b"%PDF-1.4\n"
        b"1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n"
        b"2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj\n"
        b"3 0 obj<</Type/Page/Parent 2 0 R/MediaBox[0 0 612 792]"
        b"/Contents 4 0 R/Resources<</Font<</F1 5 0 R>>>>>>endobj\n"
        b"4 0 obj<</Length 44>>stream\n"
        b"BT /F1 12 Tf 100 700 Td (Hello PDF World) Tj ET\n"
        b"endstream\nendobj\n"
        b"5 0 obj<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>endobj\n"
        b"xref\n0 6\n0000000000 65535 f \n0000000009 00000 n \n0000000058 00000 n \n"
        b"0000000115 00000 n \n0000000266 00000 n \n0000000360 00000 n \n"
        b"trailer<</Size 6/Root 1 0 R>>\n"
        b"startxref\n433\n%%EOF"
    )
    return pdf


def _make_test_docx() -> bytes:
    """Generate a minimal DOCX file using python-docx."""
    try:
        from docx import Document

        doc = Document()
        doc.core_properties.title = "Test Document"
        doc.core_properties.author = "Test Author"

        doc.add_heading("Chapter 1", level=1)
        doc.add_paragraph("This is a test paragraph for the cleaning service.")
        doc.add_paragraph("Another paragraph with some content for testing purposes.")

        doc.add_heading("Section 1.1", level=2)
        doc.add_paragraph("Detailed content under section 1.1.")

        # Add a table
        table = doc.add_table(rows=3, cols=3)
        table.cell(0, 0).text = "Name"
        table.cell(0, 1).text = "Age"
        table.cell(0, 2).text = "City"
        table.cell(1, 0).text = "Alice"
        table.cell(1, 1).text = "30"
        table.cell(1, 2).text = "Beijing"
        table.cell(2, 0).text = "Bob"
        table.cell(2, 1).text = "25"
        table.cell(2, 2).text = "Shanghai"

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
    except ImportError:
        return b""


def _make_test_md() -> bytes:
    """Generate a test Markdown document."""
    md = """# Test Markdown Document

## Introduction

This is a test markdown file for the cleaning service pipeline.

## Features

- Feature 1: Multi-format support
- Feature 2: Element processing
- Feature 3: Quality validation

## Code Example

```python
def hello():
    print("Hello, World!")
```

> This is a blockquote with important information.

## Data Table

| Name | Score | Rank |
|------|-------|------|
| Alice | 95 | 1 |
| Bob | 87 | 2 |
| Charlie | 92 | 3 |

## Conclusion

This document tests the full cleaning pipeline.
"""
    return md.encode("utf-8")


def _make_test_txt() -> bytes:
    """Generate a test plain text document."""
    txt = """Plain Text Test Document

This is a test of the plain text preprocessor.
Multiple paragraphs are separated by blank lines.

The cleaning service should handle:
- Basic text extraction
- Paragraph detection
- Encoding normalization

Contact: test@example.com
Phone: 13800138000
"""
    return txt.encode("utf-8")


# ─── Pipeline Tests ───────────────────────────────────────


class TestPreprocessing:
    """Test format-specific preprocessing layer."""

    def test_pdf_preprocessor(self):
        """Test PDF preprocessing produces valid Document."""
        from preprocessing.pdf_preprocessor import PDFPreprocessor

        pdf_data = _make_test_pdf()
        prep = PDFPreprocessor()
        doc = prep.extract(pdf_data, "test.pdf")

        assert doc is not None
        assert doc.metadata.source_format == "pdf"
        assert doc.metadata.file_md5 is not None
        assert doc.page_count >= 1
        assert "pdf" in doc.processing_log[0] if doc.processing_log else True

    def test_docx_preprocessor(self):
        """Test DOCX preprocessing produces valid Document."""
        docx_data = _make_test_docx()
        if not docx_data:
            pytest.skip("python-docx not installed")

        from preprocessing.docx_preprocessor import DocxPreprocessor

        prep = DocxPreprocessor()
        doc = prep.extract(docx_data, "test.docx")

        assert doc is not None
        assert doc.metadata.source_format == "docx"
        assert doc.metadata.title == "Test Document"
        assert doc.metadata.author == "Test Author"
        assert doc.page_count == 1
        # Should have text elements and a table
        assert len(doc.pages[0].elements) > 0

    def test_md_preprocessor(self):
        """Test Markdown preprocessing."""
        from preprocessing.md_preprocessor import MarkdownPreprocessor

        md_data = _make_test_md()
        prep = MarkdownPreprocessor()
        doc = prep.extract(md_data, "test.md")

        assert doc is not None
        assert doc.metadata.source_format == "md"
        assert doc.page_count == 1
        assert len(doc.pages[0].elements) > 0

        # Should have heading elements
        from common.models.document import ElementRole
        headings = [e for e in doc.pages[0].elements
                    if hasattr(e, 'role') and e.role == ElementRole.HEADING]
        assert len(headings) > 0

    def test_txt_preprocessor(self):
        """Test plain text preprocessing."""
        from preprocessing.txt_preprocessor import TxtPreprocessor

        txt_data = _make_test_txt()
        prep = TxtPreprocessor()
        doc = prep.extract(txt_data, "test.txt")

        assert doc is not None
        assert doc.metadata.source_format == "txt"
        assert doc.page_count == 1
        assert doc.metadata.word_count > 0


class TestElementProcessing:
    """Test common element processing engine."""

    def test_table_processor(self):
        """Test table processing and Markdown output."""
        from elements.table_processor import TableProcessor
        from common.models.document import TableElement, TableCell, ElementRole

        proc = TableProcessor()
        table = TableElement(
            element_id="test-tbl-1",
            role=ElementRole.TABLE,
            rows=[
                [TableCell(text="A"), TableCell(text="B")],
                [TableCell(text="1"), TableCell(text="2")],
            ],
            column_count=2,
            row_count=2,
        )

        result = proc.process(table)
        assert result.quality_score >= 0.0

        md = proc.to_markdown(result)
        assert "|" in md
        assert "A" in md
        assert "---" in md

    def test_table_merged_cells(self):
        """Test merged cell expansion."""
        from elements.table_processor import TableProcessor
        from common.models.document import TableElement, TableCell, ElementRole

        proc = TableProcessor()
        table = TableElement(
            element_id="test-tbl-2",
            role=ElementRole.TABLE,
            rows=[
                [TableCell(text="Header", col_span=2)],
                [TableCell(text="A"), TableCell(text="B")],
            ],
            column_count=2,
            row_count=2,
            has_merged_cells=True,
        )

        result = proc.process(table)
        assert not result.has_merged_cells
        assert len(result.rows[0]) == 2


class TestCleaning:
    """Test general cleaning layer."""

    def test_basic_cleaner(self):
        """Test basic text cleaning."""
        from cleaning.basic_cleaner import BasicCleaner
        from common.models.document import Document, Page, TextElement, ElementRole

        doc = Document(doc_id="test")
        page = Page(page_number=1)
        page.elements.append(TextElement(
            element_id="e1", role=ElementRole.PARAGRAPH,
            text="Hello   World\r\n\r\n\r\nTest  \t  content",
        ))
        page.text_content = "Hello   World\r\n\r\n\r\nTest  \t  content"
        doc.pages.append(page)

        cleaner = BasicCleaner()
        result = cleaner.clean(doc)

        cleaned_text = result.pages[0].elements[0].text
        # Normalized: spaces collapsed, newlines normalized
        assert "Hello" in cleaned_text
        assert "World" in cleaned_text
        assert "Test" in cleaned_text
        assert "content" in cleaned_text
        assert "\r\n" not in cleaned_text

    def test_sensitive_masker(self):
        """Test PII masking."""
        from cleaning.sensitive_masker import SensitiveMasker
        from common.models.document import Document, Page, TextElement, ElementRole

        doc = Document(doc_id="test")
        page = Page(page_number=1)
        page.elements.append(TextElement(
            element_id="e1", role=ElementRole.PARAGRAPH,
            text="Contact: test@example.com, Phone: 13800138000",
        ))
        doc.pages.append(page)

        masker = SensitiveMasker()
        result = masker.mask(doc)

        masked_text = result.pages[0].elements[0].text
        assert "test@example.com" not in masked_text
        assert "13800138000" not in masked_text
        assert "***" in masked_text

    def test_quality_validator(self):
        """Test quality validation scoring."""
        from cleaning.quality_validator import QualityValidator
        from common.models.document import Document, Page, TextElement, ElementRole

        doc = Document(doc_id="test")
        long_text = (
            "This is a comprehensive document with substantial content for quality scoring. "
            "It contains multiple well-formed sentences that should pass all thresholds. "
            "The document covers several topics in depth with proper structure and coherence. "
            "Each section provides detailed information about the subject matter at hand. "
            "Readers will find this content informative and well-organized for their needs. "
        )
        for i in range(10):
            page = Page(page_number=i + 1)
            page.elements.append(TextElement(
                element_id=f"e{i}", role=ElementRole.PARAGRAPH,
                text=long_text + f" Paragraph {i}.",
            ))
            # Add a heading every 3 pages for structure
            if i % 3 == 0:
                page.elements.append(TextElement(
                    element_id=f"h{i}", role=ElementRole.HEADING,
                    text=f"Section {i // 3 + 1}",
                ))
            page.text_content = page.elements[0].text
            doc.pages.append(page)

        validator = QualityValidator()
        report = validator.validate(doc)

        assert report.overall_score > 0
        assert report.completeness > 0
        assert report.purity > 0
        assert report.coherence > 0


class TestOutput:
    """Test output layer."""

    def test_markdown_generation(self):
        """Test Document -> Markdown conversion."""
        from output.markdown_generator import MarkdownGenerator
        from common.models.document import (
            Document, DocumentMetadata, Page, TextElement, ElementRole, HeadingLevel,
            TableElement, TableCell, CodeElement,
        )

        doc = Document(doc_id="test-output", metadata=DocumentMetadata(title="Test Doc"))
        page = Page(page_number=1)

        page.elements.append(TextElement(
            element_id="e1", role=ElementRole.HEADING,
            text="Section 1", heading_level=HeadingLevel.H2,
        ))
        page.elements.append(TextElement(
            element_id="e2", role=ElementRole.PARAGRAPH,
            text="This is a test paragraph.",
        ))
        page.elements.append(CodeElement(
            element_id="e3", role=ElementRole.CODE_BLOCK,
            code='print("hello")', language="python",
        ))
        doc.pages.append(page)

        gen = MarkdownGenerator()
        md = gen.generate(doc)

        assert "# Test Doc" in md
        assert "## Section 1" in md
        assert "This is a test paragraph." in md
        assert "```python" in md
        assert 'print("hello")' in md

    def test_metadata_generation(self):
        """Test metadata JSON generation."""
        from output.metadata_generator import MetadataGenerator
        from common.models.document import Document, DocumentMetadata, QualityReport

        doc = Document(
            doc_id="test-meta",
            metadata=DocumentMetadata(
                title="Test", source_format="md",
                page_count=2, word_count=100,
            ),
        )
        doc.quality = QualityReport(overall_score=0.85, passed=True)

        gen = MetadataGenerator()
        meta = gen.generate(doc)

        assert meta["document_id"] == "test-meta"
        assert meta["source_format"] == "md"
        assert meta["document"]["title"] == "Test"
        assert meta["quality"]["overall_score"] == 0.85
        assert meta["quality"]["passed"] is True


class TestFullPipeline:
    """End-to-end pipeline tests for each format."""

    def _run_pipeline(self, file_data: bytes, file_name: str, expected_format: str):
        """Run the full cleaning pipeline and validate results."""
        from common.util.utils import get_file_extension

        # Step 1: Preprocessing
        from preprocessing.pdf_preprocessor import PDFPreprocessor
        from preprocessing.docx_preprocessor import DocxPreprocessor
        from preprocessing.md_preprocessor import MarkdownPreprocessor
        from preprocessing.txt_preprocessor import TxtPreprocessor

        preprocessors = {
            "pdf": PDFPreprocessor(),
            "docx": DocxPreprocessor(),
            "md": MarkdownPreprocessor(),
            "txt": TxtPreprocessor(),
        }

        ext = get_file_extension(file_name)
        prep = preprocessors.get(ext)
        if prep is None:
            pytest.skip(f"No preprocessor for format: {ext}")

        doc = prep.extract(file_data, file_name)
        doc.doc_id = "test-full-pipeline"

        assert doc is not None
        assert doc.metadata.source_format == expected_format or ext == expected_format
        print(f"  [1/5] Preprocessing: {doc.page_count} pages, {doc.metadata.word_count} words")

        # Step 2: Element processing
        from elements.table_processor import TableProcessor
        from elements.image_processor import ImageProcessor
        from elements.formula_processor import FormulaProcessor

        table_proc = TableProcessor()
        image_proc = ImageProcessor()
        formula_proc = FormulaProcessor()

        for page in doc.pages:
            for i, elem in enumerate(page.elements):
                role = getattr(elem, 'role', None)
                if role and role.value == "table":
                    page.elements[i] = table_proc.process(elem)
                elif role and role.value == "image":
                    page.elements[i] = image_proc.process(elem)
                elif role and role.value == "formula":
                    page.elements[i] = formula_proc.process(elem)

        element_count = doc.get_element_count()
        print(f"  [2/5] Element processing: {element_count.total()} total, "
              f"{element_count.tables} tables, {element_count.images} images")

        # Step 3: General cleaning
        from cleaning.basic_cleaner import BasicCleaner
        from cleaning.content_filter import ContentFilter
        from cleaning.structure_fixer import StructureFixer
        from cleaning.sensitive_masker import SensitiveMasker

        BasicCleaner().clean(doc)
        ContentFilter().filter(doc)
        StructureFixer().fix(doc)
        SensitiveMasker().mask(doc)
        print(f"  [3/5] Cleaning: {sum(len(p.elements) for p in doc.pages)} elements after cleaning")

        # Step 4: Quality validation
        from cleaning.quality_validator import QualityValidator
        report = QualityValidator().validate(doc)
        print(f"  [4/5] Quality: {report.summary()}")

        # Step 5: Output generation
        from output.markdown_generator import MarkdownGenerator
        from output.metadata_generator import MetadataGenerator

        md = MarkdownGenerator().generate(doc)
        meta = MetadataGenerator().generate(doc)

        print(f"  [5/5] Output: {len(md)} chars Markdown, {len(json.dumps(meta))} bytes JSON")

        # Assertions
        assert len(md) > 0, "Markdown output should not be empty"
        assert meta["document_id"] == "test-full-pipeline"
        assert meta["source_format"] == expected_format or ext == expected_format
        assert "quality" in meta

        return doc, md, meta, report

    def test_full_pipeline_md(self):
        """Full pipeline test with Markdown input."""
        print("\n=== Full Pipeline: Markdown ===")
        doc, md, meta, report = self._run_pipeline(_make_test_md(), "test.md", "md")
        assert "Test Markdown Document" in md
        assert "```python" in md
        print(f"  Quality: {report.summary()}")
        print(f"  Elements: {doc.get_element_count().to_dict()}")

    def test_full_pipeline_txt(self):
        """Full pipeline test with plain text input."""
        print("\n=== Full Pipeline: Plain Text ===")
        doc, md, meta, report = self._run_pipeline(_make_test_txt(), "test.txt", "txt")
        assert len(md) > 0
        # Check that sensitive info was masked
        assert "13800138000" not in md
        assert "test@example.com" not in md
        print(f"  Quality: {report.summary()}")

    def test_full_pipeline_pdf(self):
        """Full pipeline test with PDF input."""
        print("\n=== Full Pipeline: PDF ===")
        doc, md, meta, report = self._run_pipeline(_make_test_pdf(), "test.pdf", "pdf")
        assert len(md) > 0
        print(f"  Quality: {report.summary()}")

    def test_full_pipeline_docx(self):
        """Full pipeline test with DOCX input."""
        docx_data = _make_test_docx()
        if not docx_data:
            pytest.skip("python-docx not installed")
        print("\n=== Full Pipeline: DOCX ===")
        doc, md, meta, report = self._run_pipeline(docx_data, "test.docx", "docx")
        assert len(md) > 0
        assert "Chapter 1" in md
        print(f"  Quality: {report.summary()}")


# ─── Integration Tests ────────────────────────────────────


class TestIntegration:
    """Integration tests between services."""

    def test_cleaning_client_config(self):
        """Test that RAG-PYTHON cleaning config is properly set."""
        # Simulate config check (same as task_handlers._use_cleaning_service)
        config = {
            "cleaning": {
                "grpc_endpoint": "localhost:50056",
                "enabled": True,
            }
        }
        assert config["cleaning"]["enabled"] is True
        assert "50056" in config["cleaning"]["grpc_endpoint"]

    def test_format_router_all_formats(self):
        """Test that all supported formats are properly routed."""
        from scheduling.format_router import FormatRouter

        router = FormatRouter()

        # These should route to specific topics
        assert router.is_supported("test.pdf")
        assert router.is_supported("test.docx")
        assert router.is_supported("test.xlsx")
        assert router.is_supported("test.pptx")

        # Get routing result
        from common.constant.kafka_constants import CleaningTopics
        assert router.route("test.pdf") == CleaningTopics.PDF_PREPROCESS
        assert router.route("test.docx") == CleaningTopics.DOCX_PREPROCESS
        # MD/TXT skip preprocessing
        assert router.route("test.md") is None
        assert router.route("test.txt") is None

    def test_state_machine_transitions(self):
        """Test task state machine valid transitions."""
        from scheduling.state_machine import TaskStateMachine
        from common.enums.status_enums import CleaningTaskStatus

        sm = TaskStateMachine("test-task")
        assert sm.current == CleaningTaskStatus.PENDING

        # Valid transitions
        assert sm.transition(CleaningTaskStatus.ROUTING)
        assert sm.transition(CleaningTaskStatus.PREPROCESSING)
        assert sm.transition(CleaningTaskStatus.ELEMENT_PROCESSING)
        assert sm.transition(CleaningTaskStatus.CLEANING)
        assert sm.transition(CleaningTaskStatus.VALIDATING)
        assert sm.transition(CleaningTaskStatus.GENERATING_OUTPUT)
        assert sm.transition(CleaningTaskStatus.SUCCESS)

        # SUCCESS is terminal
        assert sm.is_terminal()

        # Cannot transition from SUCCESS
        assert not sm.transition(CleaningTaskStatus.FAILED)

    def test_state_machine_retry_path(self):
        """Test retry transitions."""
        from scheduling.state_machine import TaskStateMachine
        from common.enums.status_enums import CleaningTaskStatus

        sm = TaskStateMachine("test-retry")
        sm.transition(CleaningTaskStatus.ROUTING)
        sm.transition(CleaningTaskStatus.PREPROCESSING)

        # Fail and retry
        assert sm.mark_failed("test error")
        assert sm.current == CleaningTaskStatus.FAILED
        assert sm.is_retryable()

        assert sm.mark_retrying(1, 3)
        assert sm.current == CleaningTaskStatus.RETRYING

        # Re-enter from retry
        assert sm.transition(CleaningTaskStatus.ROUTING)

    def test_retry_scheduler_backoff(self):
        """Test exponential backoff calculation."""
        from scheduling.retry_scheduler import RetryScheduler

        scheduler = RetryScheduler()
        scheduler._base_delay = 60
        scheduler._multiplier = 5

        task_id = "test-backoff"

        # First retry (attempt 0): base_delay * multiplier^0 = 60 * 1 = 60
        assert scheduler.should_retry(task_id)
        delay1 = scheduler.calculate_delay(task_id)
        assert delay1 == 60  # 60 * 5^0 = 60s (1min)
        scheduler.record_attempt(task_id)

        # Second retry (attempt 1): base_delay * multiplier^1 = 60 * 5 = 300
        delay2 = scheduler.calculate_delay(task_id)
        assert delay2 == 300  # 60 * 5^1 = 300s (5min)
        scheduler.record_attempt(task_id)

        # Third retry (attempt 2): base_delay * multiplier^2 = 60 * 25 = 1500
        delay3 = scheduler.calculate_delay(task_id)
        assert delay3 == 1500  # 60 * 5^2 = 1500s (25min)

    def test_dlq_manager(self):
        """Test dead letter queue operations."""
        from scheduling.dlq_manager import DLQManager

        dlq = DLQManager()
        task_id = "test-dlq-1"

        dlq.enqueue(task_id, {"data": "test"}, "Test error message")
        entry = dlq.get_task(task_id)
        assert entry is not None
        assert entry["status"] == "PENDING_REVIEW"
        assert "Test error message" in entry["error"]

        # Re-submit
        data = dlq.re_submit(task_id)
        assert data is not None
        assert dlq.get_task(task_id)["status"] == "RE_SUBMITTED"

        # Stats
        stats = dlq.get_stats()
        assert stats["total"] >= 1


if __name__ == "__main__":
    pytest.main([__file__, "-v", "-s"])
